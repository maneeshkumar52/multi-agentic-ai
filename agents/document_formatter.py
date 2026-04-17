import datetime
from pathlib import Path
from utils import Config, Logger

class DocumentFormatterAgent:
    def __init__(self):
        self.logger = Logger("DocumentFormatterAgent")

    def format_documents(self, compliance_result, query, filename_prefix='search_report'):
        """Generate documents from validated compliance results"""
        try:
            self.logger.log_agent_activity("DocumentFormatterAgent", f"Starting document generation for: {query}")

            if not compliance_result or not compliance_result.get('approved'):
                return {
                    'status': 'error',
                    'markdown_path': None,
                    'pdf_path': None,
                    'pdf_generated': False,
                    'docx_path': None,
                    'docx_generated': False,
                    'files_created': [],
                    'content_validated': False,
                    'error': 'Invalid or unapproved compliance result'
                }

            # Generate timestamp
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

            # File paths
            md_filename = f"{filename_prefix}_{timestamp}.md"
            pdf_filename = f"{filename_prefix}_{timestamp}.pdf"
            docx_filename = f"{filename_prefix}_{timestamp}.docx"

            md_path = Config.OUTPUT_DIR / md_filename
            pdf_path = Config.OUTPUT_DIR / pdf_filename
            docx_path = Config.OUTPUT_DIR / docx_filename

            # Generate documents
            md_success = self._generate_markdown(md_path, compliance_result, query, timestamp)
            pdf_success = self._generate_pdf(pdf_path, compliance_result, query, timestamp)
            docx_success = self._generate_word(docx_path, compliance_result, query, timestamp)

            files_created = []
            if md_success:
                files_created.append(str(md_path))
            if pdf_success:
                files_created.append(str(pdf_path))
            if docx_success:
                files_created.append(str(docx_path))

            result = {
                'status': 'success',
                'markdown_path': str(md_path) if md_success else None,
                'pdf_path': str(pdf_path) if pdf_success else None,
                'pdf_generated': pdf_success,
                'docx_path': str(docx_path) if docx_success else None,
                'docx_generated': docx_success,
                'files_created': files_created,
                'content_validated': True
            }

            self.logger.log_agent_activity("DocumentFormatterAgent",
                f"Documents generated: MD={md_success}, PDF={pdf_success}, DOCX={docx_success}")

            return result

        except Exception as e:
            self.logger.log_error(f"Document generation failed: {e}")
            return {
                'status': 'error',
                'markdown_path': None,
                'pdf_path': None,
                'pdf_generated': False,
                'docx_path': None,
                'docx_generated': False,
                'files_created': [],
                'content_validated': False,
                'error': str(e)
            }

    def _generate_markdown(self, filepath, compliance_result, query, timestamp):
        """Generate Markdown document"""
        try:
            results = compliance_result.get('cleaned_results', [])
            ai_summary = compliance_result.get('validated_summary')

            content = []
            content.append(f"# Search Report: {query}")
            content.append(f"**Generated:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            content.append("")

            if ai_summary:
                content.append("## AI Summary")
                content.append(ai_summary)
                content.append("")

            content.append("## Search Results")
            content.append("")

            for i, result in enumerate(results, 1):
                content.append(f"### {i}. {result['title']}")
                content.append(f"**URL:** {result['url']}")
                content.append(f"**Snippet:** {result['snippet']}")
                content.append("")

            content.append("## Compliance Validation")
            content.append("✓ Content validated and citation-checked")
            content.append(f"✓ Quality Score: {compliance_result.get('quality_score', 0):.2f}")
            content.append(f"✓ Results: {len(results)} validated sources")
            content.append("")

            # Write file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))

            return True

        except Exception as e:
            self.logger.log_error(f"Markdown generation failed: {e}")
            return False

    def _generate_pdf(self, filepath, compliance_result, query, timestamp):
        """Generate PDF document using PyMuPDF"""
        try:
            # Import PyMuPDF with fallback
            try:
                import fitz  # PyMuPDF < 1.24
            except ImportError:
                import pymupdf as fitz  # PyMuPDF >= 1.24

            results = compliance_result.get('cleaned_results', [])
            ai_summary = compliance_result.get('validated_summary')

            # Create PDF document
            doc = fitz.open()
            page = doc.new_page()

            # Set font
            font = "helv"  # Helvetica

            # Title
            page.insert_text((50, 50), f"Search Report: {query}", fontsize=16, fontname=font)
            y_pos = 80

            # Timestamp
            timestamp_str = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            page.insert_text((50, y_pos), f"Generated: {timestamp_str}", fontsize=10, fontname=font)
            y_pos += 30

            # AI Summary
            if ai_summary:
                page.insert_text((50, y_pos), "AI Summary:", fontsize=12, fontname=font)
                y_pos += 20
                # Split summary into lines
                words = ai_summary.split()
                line = ""
                for word in words:
                    if fitz.get_text_length(line + " " + word, fontname=font, fontsize=10) < 500:
                        line += " " + word
                    else:
                        page.insert_text((50, y_pos), line.strip(), fontsize=10, fontname=font)
                        y_pos += 15
                        line = word
                if line:
                    page.insert_text((50, y_pos), line.strip(), fontsize=10, fontname=font)
                    y_pos += 25

            # Results
            page.insert_text((50, y_pos), "Search Results:", fontsize=12, fontname=font)
            y_pos += 20

            for i, result in enumerate(results, 1):
                if y_pos > 750:  # New page if needed
                    page = doc.new_page()
                    y_pos = 50

                title = f"{i}. {result['title']}"
                page.insert_text((50, y_pos), title, fontsize=11, fontname=font)
                y_pos += 15

                url = f"URL: {result['url']}"
                page.insert_text((70, y_pos), url, fontsize=9, fontname=font)
                y_pos += 12

                # Snippet (truncated)
                snippet = result['snippet'][:200] + "..." if len(result['snippet']) > 200 else result['snippet']
                words = snippet.split()
                line = ""
                for word in words:
                    if fitz.get_text_length(line + " " + word, fontname=font, fontsize=9) < 450:
                        line += " " + word
                    else:
                        page.insert_text((70, y_pos), line.strip(), fontsize=9, fontname=font)
                        y_pos += 12
                        line = word
                if line:
                    page.insert_text((70, y_pos), line.strip(), fontsize=9, fontname=font)
                    y_pos += 20

            # Compliance badge
            if y_pos > 700:
                page = doc.new_page()
                y_pos = 50

            page.insert_text((50, y_pos), "✓ Content Validated & Citation-Checked", fontsize=10, fontname=font)
            y_pos += 15
            quality_score = compliance_result.get('quality_score', 0)
            page.insert_text((50, y_pos), f"Quality Score: {quality_score:.2f}", fontsize=10, fontname=font)

            # Save PDF
            doc.save(filepath)
            doc.close()

            return True

        except Exception as e:
            self.logger.log_error(f"PDF generation failed: {e}")
            return False

    def _generate_word(self, filepath, compliance_result, query, timestamp):
        """Generate Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            results = compliance_result.get('cleaned_results', [])
            ai_summary = compliance_result.get('validated_summary')

            doc = Document()

            # Title
            title = doc.add_heading(f'Search Report: {query}', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Timestamp
            timestamp_str = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            p = doc.add_paragraph(f'Generated: {timestamp_str}')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()

            # AI Summary
            if ai_summary:
                doc.add_heading('AI Summary', level=1)
                doc.add_paragraph(ai_summary)
                doc.add_paragraph()

            # Results
            doc.add_heading('Search Results', level=1)

            for i, result in enumerate(results, 1):
                # Title
                doc.add_heading(f'{i}. {result["title"]}', level=2)

                # URL
                p = doc.add_paragraph()
                p.add_run('URL: ').bold = True
                p.add_run(result['url'])

                # Snippet
                p = doc.add_paragraph()
                p.add_run('Snippet: ').bold = True
                p.add_run(result['snippet'])

                doc.add_paragraph()

            # Compliance validation
            doc.add_heading('Compliance Validation', level=1)
            doc.add_paragraph('✓ Content Validated & Citation-Checked')
            quality_score = compliance_result.get('quality_score', 0)
            doc.add_paragraph(f'Quality Score: {quality_score:.2f}')
            doc.add_paragraph(f'Results: {len(results)} validated sources')

            # Save document
            doc.save(filepath)

            return True

        except Exception as e:
            self.logger.log_error(f"Word generation failed: {e}")
            return False